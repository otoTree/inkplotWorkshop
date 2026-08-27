from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# The document is Chinese-first. LibreOffice on macOS does not reliably honor
# w:eastAsia fallback when the primary run font lacks CJK glyphs, so use the
# installed CJK family for all run font slots.
BASE_FONT = "Hiragino Sans GB"
EAST_ASIA_FONT = "Hiragino Sans GB"
BLACK = "000000"
BODY = "222222"
MUTED = "666666"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D8DEE6"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_east_asia_font(element, font_name: str = EAST_ASIA_FONT) -> None:
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, *, size=None, color=None, bold=None, italic=None, name=BASE_FONT):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    set_east_asia_font(run._element)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: float, color: str, bold=False, name=BASE_FONT):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: list[int]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def infer_column_widths(rows: list[list[str]]) -> list[int]:
    col_count = max(len(row) for row in rows)
    weights = []
    for col in range(col_count):
        values = [row[col] if col < len(row) else "" for row in rows]
        longest = max((len(re.sub(r"[`*_]", "", value)) for value in values), default=1)
        weights.append(max(8, min(42, longest)))

    min_width = 900 if col_count >= 5 else 1100
    available = CONTENT_WIDTH_DXA - min_width * col_count
    if available < 0:
        base = CONTENT_WIDTH_DXA // col_count
        widths = [base] * col_count
    else:
        total_weight = sum(weights)
        widths = [min_width + round(available * weight / total_weight) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_num_definition(document: Document, *, kind: str, start_value: int = 1) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), str(start_value))
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")


def add_inline(paragraph, text: str, *, base_bold=False, base_color=BODY, size=11):
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=base_color, bold=base_bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Menlo", size=max(9, size - 1), color=DARK_BLUE, bold=base_bold)
            set_east_asia_font(run._element)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=base_color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=base_color, italic=True, bold=base_bold)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=base_color, bold=base_bold)


def add_callout(document: Document, text: str):
    paragraph = document.add_paragraph(style="Callout")
    add_inline(paragraph, text, base_color=DARK_BLUE)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    p_pr.append(borders)
    return paragraph


def add_table(document: Document, rows: list[list[str]]):
    if len(rows) < 2:
        return
    col_count = max(len(row) for row in rows)
    normalized = [row + [""] * (col_count - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=col_count)
    table.style = "Table Grid"
    widths = infer_column_widths(normalized)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])

    for row_idx, row in enumerate(normalized):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                shade_cell(cell, LIGHT_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(
                paragraph,
                value,
                base_bold=row_idx == 0,
                base_color=BLACK if row_idx == 0 else BODY,
                size=9.2 if col_count >= 5 else 9.8,
            )

    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int):
    raw = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw.append(lines[index].strip())
        index += 1
    rows = []
    for line in raw:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows, index


def configure_styles(document: Document):
    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    set_style_font(title, size=23, color=BLACK, bold=True)
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    title_p_pr = title.element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, size=14, color="373737")
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.paragraph_format.keep_with_next = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Metadata" not in styles:
        metadata = styles.add_style("Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        metadata = styles["Metadata"]
    set_style_font(metadata, size=10.5, color=MUTED)
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(2)
    metadata.paragraph_format.line_spacing = 1.05

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    set_style_font(callout, size=10.5, color=DARK_BLUE)
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.10)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.15
    callout.paragraph_format.keep_together = True


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MUTED)


def configure_page(document: Document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("INKPLOT WORKSHOP · CURRENT-STATE PRD")
    set_run_font(hr, size=8.5, color="888888", bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("INTERNAL PRODUCT DOCUMENTATION  ·  ")
    set_run_font(fr, size=8.5, color="888888")
    add_page_field(fp)


def add_metadata_paragraph(document: Document, line: str):
    paragraph = document.add_paragraph(style="Metadata")
    if "：" in line:
        label, value = line.split("：", 1)
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, size=10.5, color=BLACK, bold=True)
        add_inline(paragraph, value.strip(), base_color=MUTED, size=10.5)
    else:
        add_inline(paragraph, line, base_color=MUTED, size=10.5)
    return paragraph


def build(markdown_path: Path, output_path: Path):
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_page(document)
    configure_styles(document)
    bullet_num_id = add_num_definition(document, kind="bullet")

    props = document.core_properties
    props.title = "Inkplot Workshop 产品需求文档｜现状版"
    props.subject = "基于 main@6f197b0 的 Current-State PRD"
    props.author = "Codex"
    props.keywords = "Inkplot Workshop, PRD, AI短剧, 分镜, 视频生成"

    index = 0
    front_matter = True
    paragraph_buffer: list[str] = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = document.add_paragraph()
            add_inline(paragraph, " ".join(item.strip() for item in paragraph_buffer))
            paragraph_buffer = []

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()

        if not line:
            flush_paragraph()
            index += 1
            continue

        if line.startswith("|"):
            flush_paragraph()
            rows, next_index = parse_table(lines, index)
            add_table(document, rows)
            index = next_index
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1:
                paragraph = document.add_paragraph(style="Title")
                add_inline(paragraph, text, base_bold=True, base_color=BLACK, size=23)
            elif front_matter and level == 2 and text.startswith("现状版"):
                paragraph = document.add_paragraph(style="Subtitle")
                add_inline(paragraph, text, base_color="373737", size=14)
            else:
                front_matter = False
                style_name = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}[level]
                paragraph = document.add_paragraph(style=style_name)
                add_inline(
                    paragraph,
                    text,
                    base_bold=True,
                    base_color=BLUE if level < 4 else DARK_BLUE,
                    size={2: 16, 3: 13, 4: 12}[level],
                )
            index += 1
            continue

        if line.startswith("> "):
            flush_paragraph()
            front_matter = False
            add_callout(document, line[2:].strip())
            index += 1
            continue

        if front_matter and re.match(r"^[^：]{2,12}：", line):
            flush_paragraph()
            add_metadata_paragraph(document, line)
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            paragraph = document.add_paragraph()
            apply_numbering(paragraph, bullet_num_id)
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if numbered:
            flush_paragraph()
            paragraph = document.add_paragraph()
            item_num_id = add_num_definition(
                document,
                kind="decimal",
                start_value=int(numbered.group(1)),
            )
            apply_numbering(paragraph, item_num_id)
            add_inline(paragraph, numbered.group(2))
            index += 1
            continue

        paragraph_buffer.append(line.rstrip("  "))
        index += 1

    flush_paragraph()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_current_prd_docx.py INPUT.md OUTPUT.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
